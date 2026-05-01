from __future__ import annotations

import io
import json
import mimetypes
from pathlib import Path

from .types import NormalizedDocument, UnsupportedFormatError

__all__ = ["FileAdapterFactory", "NormalizedDocument", "UnsupportedFormatError"]


_TEXT_EXTS = {".txt", ".md", ".markdown", ".log", ".rst", ".tsv"}


class FileAdapterFactory:
    """Single entry point: feed bytes + filename, get a NormalizedDocument.

    The factory owns the dispatch table: ext/mime → adapter callable. Each
    adapter is small and lazy-imports its heavy dependency (pdfplumber, docx,
    etc.) so apps that only handle CSVs don't pay the import cost.

    Usage:
        doc = FileAdapterFactory.normalize(content_bytes, filename="report.pdf")
        doc.text       # always populated
        doc.tabular    # list of dicts when applicable (CSV/XLSX/JSON rows)
        doc.metadata   # provenance + adapter-specific extras
    """

    @classmethod
    def normalize(
        cls,
        content: bytes,
        *,
        filename: str | None = None,
        mime: str | None = None,
    ) -> NormalizedDocument:
        ext = cls._detect_ext(filename, mime)
        adapter = cls._dispatch(ext)
        doc = adapter(content)
        doc.metadata.setdefault("filename", filename)
        doc.metadata.setdefault("ext", ext)
        doc.metadata.setdefault("size", len(content))
        return doc

    @staticmethod
    def supported_extensions() -> list[str]:
        return sorted(_DISPATCH.keys() | _TEXT_EXTS)

    @staticmethod
    def _detect_ext(filename: str | None, mime: str | None) -> str:
        if filename:
            ext = Path(filename).suffix.lower()
            if ext:
                return ext
        if mime:
            guessed = mimetypes.guess_extension(mime)
            if guessed:
                return guessed.lower()
        return ".txt"

    @staticmethod
    def _dispatch(ext: str):
        if ext in _TEXT_EXTS:
            return _adapter_text
        if ext in _DISPATCH:
            return _DISPATCH[ext]
        raise UnsupportedFormatError(
            f"No adapter for extension {ext!r}. "
            f"Supported: {', '.join(FileAdapterFactory.supported_extensions())}"
        )


def _decode(content: bytes) -> str:
    try:
        import chardet

        guess = chardet.detect(content) or {}
        encoding = guess.get("encoding") or "utf-8"
    except Exception:
        encoding = "utf-8"
    return content.decode(encoding, errors="replace")


def _adapter_text(content: bytes) -> NormalizedDocument:
    return NormalizedDocument(text=_decode(content), metadata={"adapter": "text"})


def _adapter_markdown(content: bytes) -> NormalizedDocument:
    raw = _decode(content)
    return NormalizedDocument(text=raw, metadata={"adapter": "markdown"})


def _adapter_html(content: bytes) -> NormalizedDocument:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text("\n", strip=True)
    title = (soup.title.string.strip() if soup.title and soup.title.string else None)
    return NormalizedDocument(
        text=text, metadata={"adapter": "html", "title": title}
    )


def _adapter_pdf(content: bytes) -> NormalizedDocument:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)
    text = "\n\n".join(pages)
    return NormalizedDocument(
        text=text,
        metadata={"adapter": "pdf", "page_count": len(pages)},
    )


def _adapter_docx(content: bytes) -> NormalizedDocument:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
    text = "\n".join(paragraphs)
    if tables:
        text += "\n\n" + "\n\n".join(
            "\n".join("\t".join(r) for r in t) for t in tables
        )
    return NormalizedDocument(
        text=text,
        metadata={
            "adapter": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        },
    )


def _adapter_csv(content: bytes) -> NormalizedDocument:
    import pandas as pd

    df = pd.read_csv(io.BytesIO(content))
    return _df_to_document(df, adapter="csv")


def _adapter_tsv(content: bytes) -> NormalizedDocument:
    import pandas as pd

    df = pd.read_csv(io.BytesIO(content), sep="\t")
    return _df_to_document(df, adapter="tsv")


def _adapter_xlsx(content: bytes) -> NormalizedDocument:
    import pandas as pd

    sheets = pd.read_excel(io.BytesIO(content), sheet_name=None)
    if not sheets:
        return NormalizedDocument(text="", metadata={"adapter": "xlsx", "sheets": []})

    parts: list[str] = []
    tabular: list[dict] = []
    for name, df in sheets.items():
        parts.append(f"# Sheet: {name}\n{df.to_csv(index=False)}")
        for record in df.fillna("").to_dict(orient="records"):
            tabular.append({"_sheet": name, **record})

    return NormalizedDocument(
        text="\n\n".join(parts),
        tabular=tabular,
        metadata={"adapter": "xlsx", "sheets": list(sheets.keys())},
    )


def _adapter_json(content: bytes) -> NormalizedDocument:
    raw = _decode(content)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return NormalizedDocument(text=raw, metadata={"adapter": "json", "parsed": False})

    tabular: list[dict] | None = None
    if isinstance(data, list) and data and all(isinstance(d, dict) for d in data):
        tabular = data
    elif isinstance(data, dict):
        for value in data.values():
            if isinstance(value, list) and value and all(isinstance(d, dict) for d in value):
                tabular = value
                break

    text = json.dumps(data, indent=2, ensure_ascii=False)
    return NormalizedDocument(
        text=text,
        tabular=tabular,
        metadata={"adapter": "json", "parsed": True},
    )


def _adapter_xml(content: bytes) -> NormalizedDocument:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml-xml")
    text = soup.get_text("\n", strip=True)
    return NormalizedDocument(text=text, metadata={"adapter": "xml"})


def _adapter_parquet(content: bytes) -> NormalizedDocument:
    import pandas as pd

    df = pd.read_parquet(io.BytesIO(content))
    return _df_to_document(df, adapter="parquet")


def _df_to_document(df, *, adapter: str) -> NormalizedDocument:
    text = df.to_csv(index=False)
    tabular = df.fillna("").to_dict(orient="records")
    return NormalizedDocument(
        text=text,
        tabular=tabular,
        metadata={
            "adapter": adapter,
            "rows": int(df.shape[0]),
            "columns": list(df.columns),
        },
    )


_DISPATCH = {
    ".csv": _adapter_csv,
    ".tsv": _adapter_tsv,
    ".xlsx": _adapter_xlsx,
    ".xls": _adapter_xlsx,
    ".json": _adapter_json,
    ".pdf": _adapter_pdf,
    ".docx": _adapter_docx,
    ".html": _adapter_html,
    ".htm": _adapter_html,
    ".xml": _adapter_xml,
    ".parquet": _adapter_parquet,
}
